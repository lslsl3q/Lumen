"""
T11 写作模式 REST API — 作品/章节/世界观设定 CRUD
"""

import asyncio
import json
import logging
import os
import re
import shutil
import uuid
from fastapi import APIRouter, HTTPException, UploadFile, File
from fastapi.responses import PlainTextResponse, StreamingResponse, FileResponse
from pydantic import BaseModel, Field

from lumen.config import ASSETS_DIR
from lumen.services.storage.writing import (
    create_project, list_projects, get_project, update_project, delete_project,
    create_chapter, list_chapters, get_chapter, update_chapter, delete_chapter, reorder_chapters,
    create_codex, list_codex, get_codex, update_codex, delete_codex, reorder_codex,
    # NEW:
    create_act, list_acts, get_act, update_act, delete_act, reorder_acts,
    create_scene, list_scenes, get_scene, update_scene, delete_scene, reorder_scenes,
    get_manuscript, get_manuscript_flat,
    create_snippet, list_snippets, get_snippet, update_snippet, delete_snippet,
    # Labels
    create_label, list_labels, update_label, delete_label, reorder_labels,
    # Plot System
    get_or_create_plot, update_plot,
    create_arc, list_arcs, update_arc, delete_arc, reorder_arcs,
    create_line, list_lines, update_line, delete_line, reorder_lines,
    create_node, list_nodes, update_node, delete_node, reorder_nodes,
    create_beat, list_beats, update_beat, delete_beat, reorder_beats,
    create_link, list_links_for_node, delete_link,
    get_plot_tree,
)
from lumen.services.storage.writing_snapshot import (
    create_snapshot, list_snapshots, get_snapshot_detail, restore_snapshot, delete_snapshot,
)

logger = logging.getLogger(__name__)
router = APIRouter()


# ── 请求模型 ──

class CreateActRequest(BaseModel):
    project_id: str
    title: str = ""
    numerate: bool = True


class UpdateActRequest(BaseModel):
    title: str | None = None
    numerate: bool | None = None


class ReorderActsRequest(BaseModel):
    project_id: str
    ordered_ids: list[str]


class CreateChapterV2Request(BaseModel):
    act_id: str
    project_id: str
    title: str = ""


class UpdateChapterV2Request(BaseModel):
    title: str | None = None
    numerate: bool | None = None
    show_number: bool | None = None
    act_id: str | None = None


class ReorderChaptersV2Request(BaseModel):
    act_id: str
    ordered_ids: list[str]


class CreateSceneRequest(BaseModel):
    chapter_id: str
    content: dict | None = None
    summary: str = ""
    subtitle: str = ""


class UpdateSceneRequest(BaseModel):
    content: dict | None = None
    summary: str | None = None
    subtitle: str | None = None
    chapter_id: str | None = None
    codex_ids: list[str] | None = None
    label_ids: list[str] | None = None


class ReorderScenesRequest(BaseModel):
    chapter_id: str
    ordered_ids: list[str]


class CreateProjectRequest(BaseModel):
    name: str
    description: str = ""
    channel_id: str = ""


class UpdateProjectRequest(BaseModel):
    name: str | None = None
    description: str | None = None
    metadata: dict | None = None


class CreateCodexRequest(BaseModel):
    project_id: str
    name: str
    type: str = "custom"
    parent_id: str | None = None
    description: dict = Field(default_factory=dict)
    aliases: list[str] = Field(default_factory=list)
    tags: list[str] = Field(default_factory=list)
    category: str | None = None


class UpdateCodexRequest(BaseModel):
    name: str | None = None
    type: str | None = None
    description: dict | None = None
    aliases: list[str] | None = None
    tags: list[str] | None = None
    category: str | None = None
    custom_fields: dict | None = None
    relations: list | None = None
    graph_entity_id: str | None = None
    enabled: int | None = None
    parent_id: str | None = None


class ReorderCodexRequest(BaseModel):
    ordered_ids: list[str]


class CreateSnapshotRequest(BaseModel):
    label: str = ""
    type: str = "manual"


# ── 作品 ──

@router.get("/projects")
async def api_list_projects():
    try:
        return await asyncio.to_thread(list_projects)
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/projects")
async def api_create_project(req: CreateProjectRequest):
    try:
        return await asyncio.to_thread(
            create_project, req.name, req.description, req.channel_id
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/projects/{project_id}")
async def api_get_project(project_id: str):
    proj = await asyncio.to_thread(get_project, project_id)
    if not proj:
        raise HTTPException(status_code=404, detail="作品不存在")
    return proj


@router.patch("/projects/{project_id}")
async def api_update_project(project_id: str, req: UpdateProjectRequest):
    updates = {k: v for k, v in req.model_dump().items() if v is not None}
    proj = await asyncio.to_thread(update_project, project_id, **updates)
    if not proj:
        raise HTTPException(status_code=404, detail="作品不存在")
    return proj


@router.delete("/projects/{project_id}")
async def api_delete_project(project_id: str):
    await asyncio.to_thread(delete_project, project_id)
    return {"status": "deleted"}


# ── 封面 ──

COVERS_DIR = os.path.join(ASSETS_DIR, "writing", "covers")
os.makedirs(COVERS_DIR, exist_ok=True)


@router.post("/projects/{project_id}/cover")
async def api_upload_cover(project_id: str, file: UploadFile = File(...)):
    if not file.content_type or not file.content_type.startswith("image/"):
        raise HTTPException(status_code=400, detail="Only image files are allowed")
    ext = file.content_type.split("/")[-1].replace("jpeg", "jpg")
    if ext not in ("jpg", "png", "gif", "webp", "bmp"):
        ext = "jpg"
    filename = f"{project_id}.{ext}"
    dest = os.path.join(COVERS_DIR, filename)
    # Remove old cover (different extension)
    for old in os.listdir(COVERS_DIR):
        if old.startswith(f"{project_id}.") and old != filename:
            os.remove(os.path.join(COVERS_DIR, old))
    with open(dest, "wb") as f:
        shutil.copyfileobj(file.file, f)
    cover_path = f"covers/{filename}"
    await asyncio.to_thread(update_project, project_id, metadata={"cover": cover_path})
    return {"cover": cover_path}


@router.get("/projects/{project_id}/cover")
async def api_get_cover(project_id: str):
    for ext in ("jpg", "png", "gif", "webp", "bmp"):
        path = os.path.join(COVERS_DIR, f"{project_id}.{ext}")
        if os.path.exists(path):
            return FileResponse(path, media_type=f"image/{ext.replace('jpg', 'jpeg')}")
    raise HTTPException(status_code=404, detail="No cover")


# ── 章节 ──

@router.get("/projects/{project_id}/chapters")
async def api_list_chapters(project_id: str):
    return await asyncio.to_thread(list_chapters, project_id)


@router.post("/chapters")
async def api_create_chapter(req: CreateChapterV2Request):
    return await asyncio.to_thread(
        create_chapter, req.act_id, req.project_id, req.title
    )


@router.get("/chapters/{chapter_id}")
async def api_get_chapter(chapter_id: str):
    ch = await asyncio.to_thread(get_chapter, chapter_id)
    if not ch:
        raise HTTPException(status_code=404, detail="章节不存在")
    return ch


@router.patch("/chapters/{chapter_id}")
async def api_update_chapter(chapter_id: str, req: UpdateChapterV2Request):
    updates = {k: v for k, v in req.model_dump().items() if v is not None}
    ch = await asyncio.to_thread(update_chapter, chapter_id, **updates)
    if not ch:
        raise HTTPException(status_code=404, detail="章节不存在")
    return ch


@router.delete("/chapters/{chapter_id}")
async def api_delete_chapter(chapter_id: str):
    await asyncio.to_thread(delete_chapter, chapter_id)
    return {"status": "deleted"}


@router.post("/chapters/reorder")
async def api_reorder_chapters(req: ReorderChaptersV2Request):
    await asyncio.to_thread(reorder_chapters, req.act_id, req.ordered_ids)
    return {"status": "reordered"}


# ── Act ──

@router.get("/projects/{project_id}/acts")
async def api_list_acts(project_id: str):
    return await asyncio.to_thread(list_acts, project_id)


@router.post("/acts")
async def api_create_act(req: CreateActRequest):
    return await asyncio.to_thread(create_act, req.project_id, req.title, req.numerate)


@router.get("/acts/{act_id}")
async def api_get_act(act_id: str):
    a = await asyncio.to_thread(get_act, act_id)
    if not a:
        raise HTTPException(status_code=404, detail="Act not found")
    return a


@router.patch("/acts/{act_id}")
async def api_update_act(act_id: str, req: UpdateActRequest):
    updates = {k: v for k, v in req.model_dump().items() if v is not None}
    a = await asyncio.to_thread(update_act, act_id, **updates)
    if not a:
        raise HTTPException(status_code=404, detail="Act not found")
    return a


@router.delete("/acts/{act_id}")
async def api_delete_act(act_id: str):
    await asyncio.to_thread(delete_act, act_id)
    return {"status": "deleted"}


@router.post("/acts/reorder")
async def api_reorder_acts(req: ReorderActsRequest):
    await asyncio.to_thread(reorder_acts, req.project_id, req.ordered_ids)
    return {"status": "reordered"}


# ── Codex (世界观设定) ──

@router.get("/projects/{project_id}/codex")
async def api_list_codex(project_id: str, type: str | None = None):
    return await asyncio.to_thread(list_codex, project_id, type)


@router.post("/codex")
async def api_create_codex(req: CreateCodexRequest):
    return await asyncio.to_thread(
        create_codex, req.project_id, req.name, req.type, req.parent_id, req.description, req.aliases, req.tags, req.category
    )


@router.get("/codex/{codex_id}")
async def api_get_codex(codex_id: str):
    s = await asyncio.to_thread(get_codex, codex_id)
    if not s:
        raise HTTPException(status_code=404, detail="Codex 条目不存在")
    return s


@router.patch("/codex/{codex_id}")
async def api_update_codex(codex_id: str, req: UpdateCodexRequest):
    updates = req.model_dump(exclude_unset=True)
    if not updates:
        s = await asyncio.to_thread(get_codex, codex_id)
        if not s:
            raise HTTPException(status_code=404, detail="Codex 条目不存在")
        return s
    s = await asyncio.to_thread(update_codex, codex_id, **updates)
    if not s:
        raise HTTPException(status_code=404, detail="Codex 条目不存在")
    return s


@router.delete("/codex/{codex_id}")
async def api_delete_codex(codex_id: str):
    await asyncio.to_thread(delete_codex, codex_id)
    return {"status": "deleted"}


@router.post("/codex/reorder")
async def api_reorder_codex(req: ReorderCodexRequest):
    await asyncio.to_thread(reorder_codex, req.ordered_ids)
    return {"status": "reordered"}


# ── Scene ──

@router.get("/chapters/{chapter_id}/scenes")
async def api_list_scenes(chapter_id: str):
    return await asyncio.to_thread(list_scenes, chapter_id)


@router.post("/scenes")
async def api_create_scene(req: CreateSceneRequest):
    return await asyncio.to_thread(create_scene, req.chapter_id, req.content, req.summary, req.subtitle)


@router.get("/scenes/{scene_id}")
async def api_get_scene(scene_id: str):
    s = await asyncio.to_thread(get_scene, scene_id)
    if not s:
        raise HTTPException(status_code=404, detail="Scene not found")
    return s


@router.patch("/scenes/{scene_id}")
async def api_update_scene(scene_id: str, req: UpdateSceneRequest):
    updates = {k: v for k, v in req.model_dump().items() if v is not None}
    s = await asyncio.to_thread(update_scene, scene_id, **updates)
    if not s:
        raise HTTPException(status_code=404, detail="Scene not found")
    return s


@router.delete("/scenes/{scene_id}")
async def api_delete_scene(scene_id: str):
    await asyncio.to_thread(delete_scene, scene_id)
    return {"status": "deleted"}


@router.post("/scenes/reorder")
async def api_reorder_scenes(req: ReorderScenesRequest):
    await asyncio.to_thread(reorder_scenes, req.chapter_id, req.ordered_ids)
    return {"status": "reordered"}


# ── 快照 ──

@router.get("/projects/{project_id}/snapshots")
async def api_list_snapshots(project_id: str, limit: int = 50):
    items = await asyncio.to_thread(list_snapshots, project_id, limit)
    return {"items": items, "total": len(items)}


@router.post("/projects/{project_id}/snapshots")
async def api_create_snapshot(project_id: str, req: CreateSnapshotRequest):
    try:
        return await asyncio.to_thread(
            create_snapshot, project_id, req.type, req.label
        )
    except ValueError as e:
        raise HTTPException(status_code=404, detail=str(e))
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/snapshots/{snapshot_id}")
async def api_get_snapshot(snapshot_id: str):
    snap = await asyncio.to_thread(get_snapshot_detail, snapshot_id)
    if not snap:
        raise HTTPException(status_code=404, detail="快照不存在")
    return snap


@router.post("/snapshots/{snapshot_id}/restore")
async def api_restore_snapshot(snapshot_id: str):
    try:
        return await asyncio.to_thread(restore_snapshot, snapshot_id)
    except ValueError as e:
        raise HTTPException(status_code=404, detail=str(e))
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@router.delete("/snapshots/{snapshot_id}")
async def api_delete_snapshot(snapshot_id: str):
    await asyncio.to_thread(delete_snapshot, snapshot_id)
    return {"status": "deleted"}


# ── 导出 ──

def _pm_json_to_plain_text(content_str: str) -> str:
    """Extract plain text from ProseMirror JSON."""
    try:
        doc = json.loads(content_str) if isinstance(content_str, str) else content_str
    except json.JSONDecodeError:
        return content_str
    texts = []
    def walk(node):
        if node.get("type") == "text" and node.get("text"):
            texts.append(node["text"])
        for child in node.get("content", []):
            walk(child)
    walk(doc)
    return "\n\n".join(texts)


def _pm_json_to_markdown(content_str: str) -> str:
    """Convert ProseMirror JSON to basic Markdown."""
    try:
        doc = json.loads(content_str) if isinstance(content_str, str) else content_str
    except json.JSONDecodeError:
        return content_str
    lines = []
    def walk(node):
        t = node.get("type", "")
        if t == "heading":
            lines.append("#" * node.get("attrs", {}).get("level", 1) + " " + _collect_text(node))
        elif t == "paragraph":
            lines.append(_collect_text(node))
        elif t == "bulletList":
            for item in node.get("content", []):
                lines.append("- " + _collect_text(item))
        elif t == "orderedList":
            for i, item in enumerate(node.get("content", []), 1):
                lines.append(f"{i}. " + _collect_text(item))
        elif t == "blockquote":
            for child in node.get("content", []):
                lines.append("> " + _collect_text(child))
        else:
            for child in node.get("content", []):
                walk(child)
    walk(doc)
    return "\n\n".join(lines)


def _collect_text(node) -> str:
    texts = []
    def walk(n):
        if n.get("type") == "text" and n.get("text"):
            texts.append(n["text"])
        for child in n.get("content", []):
            walk(child)
    walk(node)
    return "".join(texts)


def _safe_filename(name: str) -> str:
    """清理文件名：去除路径分隔符、换行、控制字符"""
    return re.sub(r'[/\\:\n\r\t\0]', '_', name).strip('. ')


@router.get("/projects/{project_id}/export")
async def api_export_project(project_id: str, format: str = "txt"):
    """Export entire project as TXT / Markdown / DOCX from new acts/chapters/scenes model."""
    manuscript = await asyncio.to_thread(get_manuscript, project_id)
    acts = manuscript.get("acts", [])
    if not acts:
        raise HTTPException(404, "No content to export")

    project = await asyncio.to_thread(get_project, project_id)
    title = project["name"] if project else "未命名"
    safe_title = _safe_filename(title)

    if format == "md":
        parts = [f"# {title}\n\n"]
        for act in acts:
            parts.append(f"## Act: {act['title']}\n\n")
            for ch in act.get("chapters", []):
                parts.append(f"### {ch['title']}\n\n")
                for sc in ch.get("scenes", []):
                    parts.append(_pm_json_to_markdown(sc.get("content", "{}")))
                    parts.append("\n\n---\n\n")
        return PlainTextResponse("".join(parts), media_type="text/markdown",
                                 headers={"Content-Disposition": f"attachment; filename={safe_title}.md"})

    elif format == "docx":
        from docx import Document
        from docx.shared import Pt
        from io import BytesIO

        doc = Document()
        doc.add_heading(title, level=0)
        for act in acts:
            doc.add_heading(f"Act: {act['title']}", level=1)
            for ch in act.get("chapters", []):
                doc.add_heading(ch["title"], level=2)
                for sc in ch.get("scenes", []):
                    text = _pm_json_to_plain_text(sc.get("content", "{}"))
                    for line in text.split("\n"):
                        line = line.strip()
                        if line:
                            p = doc.add_paragraph(line)
                            for run in p.runs:
                                run.font.size = Pt(12)
        buf = BytesIO()
        doc.save(buf)
        buf_data = buf.getvalue()
        return StreamingResponse(
            iter([buf_data]),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename={safe_title}.docx"},
        )

    else:  # txt
        parts = [f"{title}\n{'=' * len(title)}\n\n"]
        for act in acts:
            parts.append(f"Act: {act['title']}\n\n")
            for ch in act.get("chapters", []):
                parts.append(f"  {ch['title']}\n  {'-' * 20}\n\n")
                for sc in ch.get("scenes", []):
                    parts.append(_pm_json_to_plain_text(sc.get("content", "{}")))
                    parts.append("\n\n")
        return PlainTextResponse("".join(parts), media_type="text/plain",
                                 headers={"Content-Disposition": f"attachment; filename={safe_title}.txt"})


# ── Manuscript (bulk) ──

@router.get("/projects/{project_id}/manuscript")
async def api_get_manuscript(project_id: str):
    return await asyncio.to_thread(get_manuscript, project_id)


@router.get("/projects/{project_id}/manuscript-flat")
async def api_get_manuscript_flat(project_id: str):
    return await asyncio.to_thread(get_manuscript_flat, project_id)


# ── Snippets ──

class CreateSnippetRequest(BaseModel):
    project_id: str
    name: str = ""


class UpdateSnippetRequest(BaseModel):
    name: str | None = None
    content: str | None = None
    pinned: bool | None = None


@router.get("/projects/{project_id}/snippets")
async def api_list_snippets(project_id: str):
    return await asyncio.to_thread(list_snippets, project_id)


@router.post("/snippets")
async def api_create_snippet(req: CreateSnippetRequest):
    result = await asyncio.to_thread(create_snippet, req.project_id, req.name)
    if not result:
        raise HTTPException(500, "创建 Snippet 失败")
    return result


@router.get("/snippets/{snippet_id}")
async def api_get_snippet(snippet_id: str):
    result = await asyncio.to_thread(get_snippet, snippet_id)
    if not result:
        raise HTTPException(404, "Snippet 不存在")
    return result


@router.patch("/snippets/{snippet_id}")
async def api_update_snippet(snippet_id: str, req: UpdateSnippetRequest):
    fields = {k: v for k, v in req.model_dump().items() if v is not None}
    if not fields:
        return await asyncio.to_thread(get_snippet, snippet_id)
    result = await asyncio.to_thread(update_snippet, snippet_id, **fields)
    if not result:
        raise HTTPException(404, "Snippet 不存在")
    return result


@router.delete("/snippets/{snippet_id}")
async def api_delete_snippet(snippet_id: str):
    await asyncio.to_thread(delete_snippet, snippet_id)
    return {"ok": True}


# ── 标签 (Labels) ──

class CreateLabelRequest(BaseModel):
    name: str = ""
    color: str = "Gray"


class UpdateLabelRequest(BaseModel):
    name: str | None = None
    color: str | None = None


@router.get("/projects/{project_id}/labels")
async def api_list_labels(project_id: str):
    return await asyncio.to_thread(list_labels, project_id)


@router.post("/projects/{project_id}/labels")
async def api_create_label(project_id: str, req: CreateLabelRequest):
    return await asyncio.to_thread(create_label, project_id, req.name, req.color)


@router.patch("/labels/{label_id}")
async def api_update_label(label_id: str, req: UpdateLabelRequest):
    updates = {k: v for k, v in req.model_dump().items() if v is not None}
    result = await asyncio.to_thread(update_label, label_id, **updates)
    if not result:
        raise HTTPException(status_code=404, detail="标签不存在")
    return result


@router.delete("/labels/{label_id}")
async def api_delete_label(label_id: str):
    await asyncio.to_thread(delete_label, label_id)
    return {"ok": True}


@router.post("/projects/{project_id}/labels/reorder")
async def api_reorder_labels(project_id: str, req: ReorderRequest):
    await asyncio.to_thread(reorder_labels, project_id, req.ordered_ids)
    return {"ok": True}



# ── Plot System ──

class CreateArcRequest(BaseModel):
    plot_id: str
    title: str = ""
    summary: str = ""


class UpdateArcRequest(BaseModel):
    title: str | None = None
    summary: str | None = None


class ReorderArcsRequest(BaseModel):
    plot_id: str
    ordered_ids: list[str]


class CreateLineRequest(BaseModel):
    arc_id: str
    name: str = ""
    title: str = ""
    type: str = "subplot"
    color: str = "#6b7280"


class UpdateLineRequest(BaseModel):
    name: str | None = None
    title: str | None = None
    type: str | None = None
    color: str | None = None
    status: str | None = None
    summary: str | None = None


class ReorderLinesRequest(BaseModel):
    arc_id: str
    ordered_ids: list[str]


class CreateNodeRequest(BaseModel):
    line_id: str
    title: str = ""
    summary: str = ""
    purpose: str = ""
    scene_ids: list[str] | None = None
    start_ch: int | None = None
    end_ch: int | None = None


class UpdateNodeRequest(BaseModel):
    title: str | None = None
    summary: str | None = None
    purpose: str | None = None
    scene_ids: list[str] | None = None
    start_ch: int | None = None
    end_ch: int | None = None
    resolved: bool | None = None


class ReorderNodesRequest(BaseModel):
    line_id: str
    ordered_ids: list[str]


class CreateBeatRequest(BaseModel):
    node_id: str
    kind: str = "action"
    summary: str = ""
    effect: str = ""


class UpdateBeatRequest(BaseModel):
    kind: str | None = None
    summary: str | None = None
    effect: str | None = None


class ReorderBeatsRequest(BaseModel):
    node_id: str
    ordered_ids: list[str]


class CreateLinkRequest(BaseModel):
    source_node_id: str
    target_node_id: str
    relation: str = "trigger"
    note: str = ""


# L1: Plot — get-or-create per project

@router.get("/projects/{project_id}/plot")
async def api_get_plot(project_id: str):
    return await asyncio.to_thread(get_or_create_plot, project_id)


@router.patch("/plots/{plot_id}")
async def api_update_plot(plot_id: str, req: UpdateArcRequest):
    result = await asyncio.to_thread(update_plot, plot_id, **req.model_dump(exclude_none=True))
    if not result:
        raise HTTPException(status_code=404, detail="Plot not found")
    return result


@router.get("/projects/{project_id}/plot-tree")
async def api_get_plot_tree(project_id: str):
    return await asyncio.to_thread(get_plot_tree, project_id)


# L2: Arcs

@router.post("/plots/{plot_id}/arcs")
async def api_create_arc(plot_id: str, req: CreateArcRequest):
    req.plot_id = plot_id
    return await asyncio.to_thread(create_arc, req.plot_id, req.title, req.summary)


@router.get("/plots/{plot_id}/arcs")
async def api_list_arcs(plot_id: str):
    return await asyncio.to_thread(list_arcs, plot_id)


@router.patch("/arcs/{arc_id}")
async def api_update_arc(arc_id: str, req: UpdateArcRequest):
    result = await asyncio.to_thread(update_arc, arc_id, **req.model_dump(exclude_none=True))
    if not result:
        raise HTTPException(status_code=404, detail="Arc not found")
    return result


@router.delete("/arcs/{arc_id}")
async def api_delete_arc(arc_id: str):
    await asyncio.to_thread(delete_arc, arc_id)
    return {"status": "deleted"}


@router.post("/plots/{plot_id}/arcs/reorder")
async def api_reorder_arcs(req: ReorderArcsRequest):
    await asyncio.to_thread(reorder_arcs, req.plot_id, req.ordered_ids)
    return {"status": "reordered"}


# L3: Lines

@router.post("/arcs/{arc_id}/lines")
async def api_create_line(arc_id: str, req: CreateLineRequest):
    req.arc_id = arc_id
    return await asyncio.to_thread(create_line, req.arc_id, req.name, req.title, req.type, req.color)


@router.get("/arcs/{arc_id}/lines")
async def api_list_lines(arc_id: str):
    return await asyncio.to_thread(list_lines, arc_id)


@router.patch("/lines/{line_id}")
async def api_update_line(line_id: str, req: UpdateLineRequest):
    result = await asyncio.to_thread(update_line, line_id, **req.model_dump(exclude_none=True))
    if not result:
        raise HTTPException(status_code=404, detail="Line not found")
    return result


@router.delete("/lines/{line_id}")
async def api_delete_line(line_id: str):
    await asyncio.to_thread(delete_line, line_id)
    return {"status": "deleted"}


@router.post("/arcs/{arc_id}/lines/reorder")
async def api_reorder_lines(req: ReorderLinesRequest):
    await asyncio.to_thread(reorder_lines, req.arc_id, req.ordered_ids)
    return {"status": "reordered"}


# L4: Nodes

@router.post("/lines/{line_id}/nodes")
async def api_create_node(line_id: str, req: CreateNodeRequest):
    req.line_id = line_id
    return await asyncio.to_thread(create_node, req.line_id, req.title, req.summary, req.purpose, req.scene_ids, req.start_ch, req.end_ch)


@router.get("/lines/{line_id}/nodes")
async def api_list_nodes(line_id: str):
    return await asyncio.to_thread(list_nodes, line_id)


@router.patch("/nodes/{node_id}")
async def api_update_node(node_id: str, req: UpdateNodeRequest):
    updates = req.model_dump(exclude_none=True)
    result = await asyncio.to_thread(update_node, node_id, **updates)
    if not result:
        raise HTTPException(status_code=404, detail="Node not found")
    return result


@router.delete("/nodes/{node_id}")
async def api_delete_node(node_id: str):
    await asyncio.to_thread(delete_node, node_id)
    return {"status": "deleted"}


@router.post("/lines/{line_id}/nodes/reorder")
async def api_reorder_nodes(req: ReorderNodesRequest):
    await asyncio.to_thread(reorder_nodes, req.line_id, req.ordered_ids)
    return {"status": "reordered"}


# L5: Beats

@router.post("/nodes/{node_id}/beats")
async def api_create_beat(node_id: str, req: CreateBeatRequest):
    req.node_id = node_id
    return await asyncio.to_thread(create_beat, req.node_id, req.kind, req.summary, req.effect)


@router.get("/nodes/{node_id}/beats")
async def api_list_beats(node_id: str):
    return await asyncio.to_thread(list_beats, node_id)


@router.patch("/beats/{beat_id}")
async def api_update_beat(beat_id: str, req: UpdateBeatRequest):
    result = await asyncio.to_thread(update_beat, beat_id, **req.model_dump(exclude_none=True))
    if not result:
        raise HTTPException(status_code=404, detail="Beat not found")
    return result


@router.delete("/beats/{beat_id}")
async def api_delete_beat(beat_id: str):
    await asyncio.to_thread(delete_beat, beat_id)
    return {"status": "deleted"}


@router.post("/nodes/{node_id}/beats/reorder")
async def api_reorder_beats(req: ReorderBeatsRequest):
    await asyncio.to_thread(reorder_beats, req.node_id, req.ordered_ids)
    return {"status": "reordered"}


# PlotLink

@router.post("/plot-links")
async def api_create_link(req: CreateLinkRequest):
    return await asyncio.to_thread(create_link, req.source_node_id, req.target_node_id, req.relation, req.note)


@router.get("/nodes/{node_id}/links")
async def api_list_links(node_id: str):
    return await asyncio.to_thread(list_links_for_node, node_id)


@router.delete("/plot-links/{link_id}")
async def api_delete_link(link_id: str):
    await asyncio.to_thread(delete_link, link_id)
    return {"status": "deleted"}
